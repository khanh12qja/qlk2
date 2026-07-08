from pathlib import Path
import re

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(r"G:\codekeyen\qlk2")
INPUT = ROOT / "docs" / "phan-tich-nghiep-vu-erp-tong-the.md"
OUT_DIR = ROOT / "outputs" / "business-analysis"
OUT = OUT_DIR / "phan-tich-nghiep-vu-erp-tong-the.docx"


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_width(cell, width_in):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn("w:tcW"))
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(int(width_in * 1440)))
    tc_w.set(qn("w:type"), "dxa")


def style_table(table):
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = True
    for row_idx, row in enumerate(table.rows):
        for cell in row.cells:
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            for p in cell.paragraphs:
                p.paragraph_format.space_after = Pt(2)
                for run in p.runs:
                    run.font.name = "Calibri"
                    run.font.size = Pt(9)
            if row_idx == 0:
                set_cell_shading(cell, "E8EEF5")
                for p in cell.paragraphs:
                    for run in p.runs:
                        run.bold = True
                        run.font.color.rgb = RGBColor(31, 77, 120)


def setup_styles(doc):
    section = doc.sections[0]
    section.top_margin = Inches(0.8)
    section.bottom_margin = Inches(0.8)
    section.left_margin = Inches(0.8)
    section.right_margin = Inches(0.8)
    section.header_distance = Inches(0.45)
    section.footer_distance = Inches(0.45)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.1

    for name, size, color in [
        ("Heading 1", 18, RGBColor(11, 37, 69)),
        ("Heading 2", 15, RGBColor(46, 116, 181)),
        ("Heading 3", 12, RGBColor(31, 77, 120)),
    ]:
        style = styles[name]
        style.font.name = "Calibri"
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = color
        style.paragraph_format.space_before = Pt(12)
        style.paragraph_format.space_after = Pt(6)


def add_cover(doc):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("PHÂN TÍCH NGHIỆP VỤ ERP TỔNG THỂ")
    run.font.name = "Calibri"
    run.font.size = Pt(22)
    run.font.bold = True
    run.font.color.rgb = RGBColor(11, 37, 69)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Bản chia theo từng phần để phục vụ phân tích thiết kế trước phát triển")
    run.font.name = "Calibri"
    run.font.size = Pt(12)
    run.font.color.rgb = RGBColor(31, 77, 120)

    doc.add_paragraph()
    box = doc.add_table(rows=1, cols=2)
    box.style = "Table Grid"
    box.rows[0].cells[0].text = "Phạm vi"
    box.rows[0].cells[1].text = "Khách hàng, đơn hàng, kế toán, công việc, BOM/kho, chấm công/lương, phân quyền và workflow cấu hình."
    style_table(box)
    doc.add_page_break()


def add_toc(doc, sections):
    doc.add_heading("Mục lục nội dung", level=1)
    for title in sections:
        p = doc.add_paragraph(style="List Number")
        p.add_run(title)
    doc.add_page_break()


def clean_inline(text):
    text = re.sub(r"\*\*(.*?)\*\*", r"\1", text)
    text = text.replace("`", "")
    return text


def split_table_row(line):
    raw = line.strip().strip("|")
    return [clean_inline(cell.strip()) for cell in raw.split("|")]


def is_separator(line):
    return bool(re.match(r"^\s*\|?\s*:?-{3,}:?\s*(\|\s*:?-{3,}:?\s*)+\|?\s*$", line))


def add_table_from_rows(doc, rows):
    if not rows:
        return
    table = doc.add_table(rows=1, cols=len(rows[0]))
    table.style = "Table Grid"
    for idx, value in enumerate(rows[0]):
        table.rows[0].cells[idx].text = value
    for row in rows[1:]:
        cells = table.add_row().cells
        for idx, value in enumerate(row[: len(cells)]):
            cells[idx].text = value
    style_table(table)


def build_doc():
    OUT_DIR.mkdir(parents=True, exist_ok=True)
    source = INPUT.read_text(encoding="utf-8")
    lines = source.splitlines()
    section_titles = [line[3:].strip() for line in lines if line.startswith("## ")]

    doc = Document()
    setup_styles(doc)
    add_cover(doc)
    add_toc(doc, section_titles)

    in_code = False
    table_rows = []
    bullet_buffer = []

    def flush_table():
        nonlocal table_rows
        if table_rows:
            add_table_from_rows(doc, table_rows)
            table_rows = []

    def flush_bullets():
        nonlocal bullet_buffer
        for item in bullet_buffer:
            p = doc.add_paragraph(style="List Bullet")
            p.add_run(clean_inline(item))
        bullet_buffer = []

    for line in lines:
        stripped = line.strip()

        if stripped.startswith("```"):
            flush_table()
            flush_bullets()
            in_code = not in_code
            continue

        if in_code:
            if stripped:
                p = doc.add_paragraph()
                run = p.add_run(stripped)
                run.font.name = "Consolas"
                run.font.size = Pt(9)
            continue

        if stripped.startswith("|") and "|" in stripped[1:]:
            flush_bullets()
            if is_separator(stripped):
                continue
            table_rows.append(split_table_row(stripped))
            continue
        else:
            flush_table()

        if not stripped:
            flush_bullets()
            continue

        if stripped.startswith("# "):
            flush_bullets()
            doc.add_heading(clean_inline(stripped[2:]), level=1)
        elif stripped.startswith("## "):
            flush_bullets()
            if doc.paragraphs and doc.paragraphs[-1].text:
                doc.add_page_break()
            doc.add_heading(clean_inline(stripped[3:]), level=1)
        elif stripped.startswith("### "):
            flush_bullets()
            doc.add_heading(clean_inline(stripped[4:]), level=2)
        elif stripped.startswith("- "):
            bullet_buffer.append(stripped[2:])
        else:
            flush_bullets()
            doc.add_paragraph(clean_inline(stripped))

    flush_table()
    flush_bullets()

    footer = doc.sections[0].footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer.add_run("Tài liệu phân tích nghiệp vụ ERP tổng thể").font.size = Pt(9)

    doc.save(OUT)
    print(OUT)


if __name__ == "__main__":
    build_doc()
