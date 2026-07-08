from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


OUT_DIR = Path(r"G:\codekeyen\qlk2\outputs\qa-bom-report")
OUT_PATH = OUT_DIR / "phan-tich-qa-logic-bom-ky-thuat-xuat-kho.docx"


RISKS = [
    ["Critical", "Ton kha dung = ton vat ly - ton da giu", "Chua bat buoc moi API tra ton dung ton kha dung", "Don sau co the thay ca hang da giu va bi ban trung", "Don A giu 3 thanh, don B van thay 10 thanh va tiep tuc giu", "Moi API tra ton, check cong thuc, goi y vat tu phai dung available stock", "Tao service ton kho trung tam, cam module khac tu tinh ton rieng"],
    ["Critical", "Giu kho chi tiet sau khi chot BOM", "Chua co yeu cau transaction khi giu nhieu dong BOM", "Giu duoc dong nay, loi dong khac, don bi giu nua voi", "BOM co 12 dong, giu thanh thanh cong nhung giu phu kien loi", "Giu kho trong transaction, loi thi rollback toan bo", "Dung reservation aggregate voi trang thai pending -> committed"],
    ["Critical", "Kho xuat theo BOM thuc te", "Chua noi ro xuat kho cung phai transaction", "Da tru phu kien nhung chua tru thanh, du lieu lech", "Xuat 1 bo, tru GC roi loi khi cat TDVC", "Xuat kho atomic: cap nhat stock, reservation, movement, leftover trong 1 transaction", "Dung stock issue document, chi post khi tat ca dong hop le"],
    ["High", "Ky thuat co the sua BOM", "Chua co quy tac khoa BOM sau khi sep duyet/kho dang xu ly", "BOM bi sua trong luc kho xuat", "Ky thuat sua so luong GC khi kho da in phieu xuat", "Them trang thai BOM: draft, technical_approved, manager_approved, issued_locked", "Chi cho sua bang revision moi, khong sua truc tiep BOM da duyet"],
    ["High", "BOM thuc te la can cu xuat kho", "Chua noi ro snapshot day du cac truong", "Cong thuc/vat tu thay doi lam don cu kho doi chieu", "Doi ten vat tu, don cu hien ten moi", "Snapshot ma vat tu, ten vat tu, don vi, mau, kinh, chieu dai, so luong, nguon goi y, nguoi chot", "Luu version BOM va version cong thuc tai thoi diem sinh"],
    ["High", "Cat thanh tao phan thua", "Chua ro chon cach cap nhat thanh goc hay tao thanh moi", "Double count hoac mat ton phan du", "Thanh 2400 cat 2100, vua de thanh goc 300 vua tao thanh moi 300", "Chon 1 quy uoc: thanh goc chuyen issued, tao leftover bar moi 300", "Luu movement lineage de truy vet thanh cha/con"],
    ["High", "Phan thua >= nguong thi tao thanh moi", "Chua co nguong theo nhom vat tu/he vat tu", "Tao nhieu doan qua ngan gay rac kho", "Du 50mm van tao ton moi", "Them minReusableLength theo nhom vat tu hoac setting", "Them ly do scrap va bao cao hao hut"],
    ["High", "Giu thanh theo stockBarId", "Chua xu ly truong hop 1 thanh co the cat nhieu doan cho nhieu don", "Khoa ca thanh lam giam kha dung, hoac chia sai phan du", "Thanh 2400 co the cat 1000 va 800, nhung he thong khoa ca thanh", "Giai doan dau nen khoa ca thanh khi reserve; khi xuat cat moi tao leftover", "Sau nay ho tro nesting/multi-cut reservation"],
    ["High", "Don khac tra kho tru hang da giu", "Chua co expiry/timeout cho giu kho", "Hang bi giu vo thoi han lam sai ton kha dung", "Don treo 3 thang van giu 5 bo", "Them reservationExpiresAt va quy trinh gia han/nhac nhan/nha giu", "Dashboard hang dang giu qua han"],
    ["High", "Da coc co the giu mem", "Chua dinh nghia giu mem khac giu cung", "Nhan vien hieu nham la hang da khoa that", "Don da coc nhung chua chot BOM van bao kho da giu", "Giu mem chi la uu tien/flag, khong tru ton kha dung", "Doi ten thanh priority_hold, khong goi la reservation"],
    ["Medium", "Sep duyet roi chon cho xuat/cho san xuat", "Chua co case thieu mot phan", "Kho khong biet xuat phan duoc hay phai cho", "Thieu BPK nhung du tat ca thanh", "Them ket qua: du xuat, thieu mot phan, khong du xuat", "Cho phep tach dot giao neu nghiep vu can"],
    ["Medium", "Ky thuat doi ma phu kien", "Chua bat buoc ly do khi doi khac goi y", "Kho truy vet kho, tranh cai sau nay", "He thong goi y GC-A, ky thuat doi GC-B", "Bat buoc note khi replace/add/remove so voi BOM de xuat", "Luu diff giua BOM de xuat va BOM thuc te"],
    ["Medium", "Bien the chi khi cau tao thay doi", "Chua noi ro mau/kinh/kieu mo khong tao bien the cong thuc", "Co nguy co nhan ban cong thuc theo mau/kinh", "PT995-MB-8, PT995-PSS-10 bi tao tran lan", "Quy uoc ro: mau/kinh la filter, khong phai variant cau tao", "UI tach bo loc ton kho va bien the cau tao"],
    ["Medium", "Kho xuat theo ma vat tu cu the", "Chua noi ro thay the vat tu tai kho co can ky thuat duyet lai khong", "Kho tu doi ma lam sai ky thuat", "Kho thay TDVC10 bang TDVC12 vi con hang", "Neu kho thay ma vat tu phai tao request doi phuong an hoac can quyen override", "Kho chi de xuat thay the, ky thuat/quan ly duyet"],
    ["Medium", "Phu kien giu theo so luong", "Chua noi ro so luong am va movement type", "Ton co the am neu xuat vuot", "Ton 5, hai nguoi cung xuat 4", "Atomic decrement voi dieu kien quantityAvailable >= required", "Dung ledger + materialized balance co version"],
    ["Medium", "Phan quyen chua duoc mo ta", "Chua tach quyen NVKD, ky thuat, sep, kho", "Ai cung sua/xuat/duyet duoc", "NVKD sua BOM thuc te sau khi ky thuat chot", "RBAC: sales create, technician approve BOM, manager approve, warehouse issue", "Audit log bat buoc cho moi hanh dong quan trong"],
    ["Medium", "Workflow co hoan tat", "Chua co huy don, doi phuong an, tra hang", "Don bi ket neu khach huy hoac doi kich thuoc", "Don da giu kho nhung khach doi mau", "Them trang thai: tam dung, huy, cho ky thuat sua lai, da tra hang", "Workflow nen co transition matrix"],
    ["Medium", "Cho san xuat", "Chua ro co giu ton mot phan hay khong", "Ton kha dung sai voi don san xuat thieu hang", "Don thieu BPK nhung da giu thanh", "Cho phep giu cac dong du hoac nha toan bo theo cau hinh", "Them policy reserve_partial true/false"],
    ["Low", "Bao cao ton kho", "Chua noi ro bao cao ton vat ly, da giu, kha dung", "Bao cao kho gay nham lan", "Ke toan thay ton 10 nhung kho bao chi dung 7", "Moi bao cao ton nen co 3 cot: physical, reserved, available", "Them drill-down danh sach don dang giu"],
    ["Low", "Nguon phat sinh phan thua", "Chua noi ro barcode/ma thanh con", "Kho kho nhan dien doan thua", "Nhieu doan 300mm cung ma vat tu", "Sinh ma stockBarId moi, in tem neu can", "Sau nay them barcode/QR"],
]

TESTS = [
    ["Critical", "Giu kho", "Test 2 nguoi cung giu 1 vat tu cung luc", "Oversell", "Ton 5, 2 don cung giu 4", "Test concurrency atomic reservation", "DB transaction va optimistic locking"],
    ["Critical", "Xuat kho", "Test loi giua qua trinh xuat", "Ton lech", "Tru phu kien xong loi cat thanh", "Test rollback transaction", "Issue document post atomic"],
    ["High", "Cat thanh", "Test cat thanh co phan thua", "Mat hoac trung ton phan thua", "2400 cat 2100 con 300", "Thanh goc issued, thanh con 300 available/scrap", "Luu lineage parent-child"],
    ["High", "Sua BOM", "Test sua BOM sau khi da duyet", "Xuat sai BOM", "Kho dang xu ly, ky thuat sua so luong", "Khoa BOM hoac tao revision", "Transition matrix + revision"],
    ["High", "Ton kha dung", "Test don sau tra ton bi tru phan da giu", "Ban trung hang", "Don A giu 3, don B phai thay 7", "Tat ca API ton dung available", "Central InventoryAvailabilityService"],
    ["Medium", "Role", "Test NVKD khong duoc chot BOM/xuat kho", "Sai phan quyen", "NVKD bam xuat kho", "RBAC + audit", "Permission theo transition"],
    ["Medium", "Huy don", "Test huy don da giu kho", "Hang khong quay lai kha dung", "Huy don A, ton van bi tru", "Nha giu tu dong khi huy", "Reservation state machine"],
    ["Medium", "Doi phuong an", "Test ky thuat doi ma vat tu", "Khong truy vet duoc", "Doi GC-A sang GC-B", "Luu diff va ly do", "Approval neu doi sau khi sep duyet"],
    ["Medium", "Partial shortage", "Test thieu mot phan BOM", "He thong chi bao thieu chung", "Du 10 dong, thieu 2 dong", "Bao theo tung dong va tong hop", "Policy reserve_partial"],
    ["Low", "Bao cao", "Test bao cao ton vat ly/giu/kha dung", "So lieu kho gay nham lan", "Physical 10, reserved 3, available 7", "Hien 3 cot rieng", "Drill-down don dang giu"],
]


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_width(cell, width_in):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn("w:tcW"))
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(int(width_in * 1440)))
    tc_w.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths):
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    tbl = table._tbl
    tbl_pr = tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(sum(int(w * 1440) for w in widths)))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_grid = tbl.tblGrid
    if tbl_grid is not None:
        tbl.remove(tbl_grid)
    tbl_grid = OxmlElement("w:tblGrid")
    for width in widths:
        grid_col = OxmlElement("w:gridCol")
        grid_col.set(qn("w:w"), str(int(width * 1440)))
        tbl_grid.append(grid_col)
    tbl.insert(0, tbl_grid)
    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            set_cell_width(cell, widths[idx])
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER


def style_table(table, widths, header_fill="E8EEF5"):
    set_table_geometry(table, widths)
    for idx, row in enumerate(table.rows):
        for cell in row.cells:
            for p in cell.paragraphs:
                p.paragraph_format.space_after = Pt(2)
                for run in p.runs:
                    run.font.name = "Calibri"
                    run.font.size = Pt(8.3)
            if idx == 0:
                set_cell_shading(cell, header_fill)
                for p in cell.paragraphs:
                    for run in p.runs:
                        run.bold = True
                        run.font.color.rgb = RGBColor(31, 77, 120)


def add_heading(doc, text, level):
    p = doc.add_heading(text, level=level)
    return p


def add_bullets(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        p.add_run(item)


def add_risk_table(doc, title, rows):
    add_heading(doc, title, 2)
    headers = ["Logic hien tai", "Van de", "Rui ro", "Vi du xay ra", "Cach sua", "Co phuong an tot hon khong?"]
    widths = [1.05, 1.05, 1.05, 1.0, 1.15, 1.2]
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    for idx, h in enumerate(headers):
        table.rows[0].cells[idx].text = h
    for row in rows:
        cells = table.add_row().cells
        for idx, value in enumerate(row[1:]):
            cells[idx].text = value
    style_table(table, widths)


def setup_styles(doc):
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.1

    for name, size, color in [
        ("Heading 1", 16, RGBColor(46, 116, 181)),
        ("Heading 2", 13, RGBColor(46, 116, 181)),
        ("Heading 3", 12, RGBColor(31, 77, 120)),
    ]:
        style = styles[name]
        style.font.name = "Calibri"
        style.font.size = Pt(size)
        style.font.color.rgb = color
        style.font.bold = True
        style.paragraph_format.space_before = Pt(10)
        style.paragraph_format.space_after = Pt(6)


def add_cover(doc):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Phan tich QA Architect")
    run.font.name = "Calibri"
    run.font.size = Pt(24)
    run.font.bold = True
    run.font.color.rgb = RGBColor(11, 37, 69)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Logic BOM ky thuat, giu kho va xuat kho thuc te")
    run.font.name = "Calibri"
    run.font.size = Pt(14)
    run.font.color.rgb = RGBColor(31, 77, 120)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Goc nhin: Senior QA Architect - ERP / WMS")
    run.italic = True
    run.font.size = Pt(11)

    doc.add_paragraph()


def build_doc():
    OUT_DIR.mkdir(parents=True, exist_ok=True)
    doc = Document()
    setup_styles(doc)
    add_cover(doc)

    add_heading(doc, "1. Ket luan ngan gon", 1)
    doc.add_paragraph(
        "Logic hien tai dung huong va phu hop thuc te ERP/WMS: cong thuc chi sinh BOM de xuat, "
        "ky thuat chot BOM thuc te, kho xuat theo ma vat tu cu the va ton kha dung phai tru hang da giu."
    )
    doc.add_paragraph(
        "Tuy nhien, neu dua vao van hanh that, can bo sung cac diem bat buoc ve transaction, khoa dong thoi, "
        "snapshot BOM, phan quyen, audit trail, expiry cho giu kho va lineage cho phan thua sau cat."
    )

    add_heading(doc, "2. Logic nghiep vu nen giu", 1)
    add_bullets(
        doc,
        [
            "BOM de xuat sinh tu cong thuc, khong phai lenh xuat kho.",
            "Ky thuat co quyen them, bot, thay ma vat tu va chot BOM thuc te.",
            "Kho chi xuat theo BOM thuc te da chot.",
            "Ton kha dung = ton vat ly - ton dang giu.",
            "Thanh cat thua phai tao ton phan thua neu con du tieu chuan su dung.",
            "Don cu phai luu snapshot BOM thuc te de khong bi anh huong khi cong thuc thay doi.",
        ],
    )

    add_heading(doc, "3. Nhom rui ro va van de can xu ly", 1)
    for severity in ["Critical", "High", "Medium", "Low"]:
        add_risk_table(doc, f"{severity} findings", [row for row in RISKS if row[0] == severity])

    add_heading(doc, "4. Danh gia database, API va workflow", 1)
    table = doc.add_table(rows=1, cols=6)
    table.style = "Table Grid"
    headers = ["Logic hien tai", "Van de", "Rui ro", "Vi du xay ra", "Cach sua", "Co phuong an tot hon khong?"]
    for idx, h in enumerate(headers):
        table.rows[0].cells[idx].text = h
    db_rows = [
        ["Stock movements va stock bars", "Chua co reservation aggregate rieng", "Kho tinh ton kha dung phuc tap va de lech", "Vua co movement, vua co giu kho rieng khong dong bo", "Them bang/collection stock_reservations", "Co: tach inventory ledger, reservation va stock issue"],
        ["API tra ton", "Neu moi noi tu tinh ton se sai", "Sai ton kha dung", "API A tru reserved, API B khong tru", "Tao InventoryAvailabilityService dung chung", "Co: moi API goi qua service trung tam"],
        ["Workflow", "Chua co transition matrix", "Trang thai nhay sai", "Don da xuat van quay lai sua BOM", "Dinh nghia transition hop le theo role", "Co: state machine ro rang"],
        ["Audit trail", "Chua bat buoc log thay doi BOM/giu/xuat", "Kho truy vet khi co tranh chap", "Ky thuat doi ma vat tu khong ai biet", "Audit log bat buoc voi hanh dong quan trong", "Co: luu diff truoc/sau"],
    ]
    for row in db_rows:
        cells = table.add_row().cells
        for idx, value in enumerate(row):
            cells[idx].text = value
    style_table(table, [1.05, 1.05, 1.05, 1.0, 1.15, 1.2], header_fill="F2F4F7")

    add_heading(doc, "5. Test case can co", 1)
    add_risk_table(doc, "Danh sach test case uu tien", TESTS)

    add_heading(doc, "6. De xuat kien truc tot hon", 1)
    add_bullets(
        doc,
        [
            "Tao InventoryAvailabilityService trung tam de tinh ton vat ly, ton da giu va ton kha dung.",
            "Tao Reservation aggregate rieng, co trang thai pending, committed, released, issued, cancelled.",
            "Moi thao tac giu kho va xuat kho phai chay trong DB transaction.",
            "BOM thuc te phai co revision va snapshot; sau khi duyet thi khong sua truc tiep.",
            "Cac thay doi BOM sau duyet phai tao revision moi va can ly do.",
            "Moi phan thua sau cat phai co parent stockBarId de truy vet nguon.",
            "Phan quyen theo transition: sales tao don, technician chot BOM, manager duyet, warehouse xuat.",
        ],
    )

    add_heading(doc, "7. Ket luan QA", 1)
    doc.add_paragraph(
        "Co the trien khai logic nay, nhung khong nen dua vao van hanh neu chua co co che tru ton kha dung, "
        "reservation chi tiet, transaction rollback va khoa BOM sau duyet. Day la cac diem quyet dinh de tranh ban trung hang, "
        "lech ton kho va kho truy vet don cu."
    )

    doc.save(OUT_PATH)


if __name__ == "__main__":
    build_doc()
    print(OUT_PATH)
